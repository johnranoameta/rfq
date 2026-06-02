"""Pure-Python extraction for .docx (no Word installation required)."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def _paragraph_dict(para, index: int) -> dict[str, Any]:
    style_name = para.style.name if para.style else ""
    return {
        "index": index,
        "text": para.text,
        "style": style_name,
    }


def _table_dict(table, index: int) -> dict[str, Any]:
    rows_data: list[list[dict[str, Any]]] = []
    for r_idx, row in enumerate(table.rows, start=1):
        row_cells: list[dict[str, Any]] = []
        for c_idx, cell in enumerate(row.cells, start=1):
            row_cells.append({"row": r_idx, "col": c_idx, "text": cell.text})
        rows_data.append(row_cells)
    return {
        "index": index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "cells": rows_data,
    }


def _core_properties(doc: Document) -> dict[str, Any]:
    cp = doc.core_properties
    fields = (
        "author",
        "category",
        "comments",
        "content_status",
        "created",
        "identifier",
        "keywords",
        "language",
        "last_modified_by",
        "last_printed",
        "modified",
        "revision",
        "subject",
        "title",
        "version",
    )
    props: dict[str, Any] = {}
    for name in fields:
        val = getattr(cp, name, None)
        if isinstance(val, datetime):
            val = val.astimezone(timezone.utc).isoformat()
        if val is not None:
            props[name] = val
    return props


def _embedded_parts(doc: Document, output_dir: Path) -> list[dict[str, Any]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    embedded: list[dict[str, Any]] = []
    seen: set[str] = set()

    for rel in doc.part.rels.values():
        if rel.reltype not in (RT.EMBEDDED_PACKAGE, RT.OLE_OBJECT):
            continue
        target = rel.target_ref
        if target in seen:
            continue
        seen.add(target)
        try:
            blob = rel.target_part.blob
        except Exception:
            continue
        name = Path(target).name or f"embedded_{len(embedded) + 1}"
        out_path = output_dir / name
        out_path.write_bytes(blob)
        embedded.append(
            {
                "relationship": rel.reltype,
                "target": target,
                "path": str(out_path),
                "size_bytes": len(blob),
                "method": "docx_package",
            }
        )
    return embedded


def extract_docx(doc_path: Path, output_dir: Path) -> dict[str, Any]:
    doc_path = doc_path.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(str(doc_path))

    paragraphs = [_paragraph_dict(p, i) for i, p in enumerate(doc.paragraphs, start=1)]
    tables = [_table_dict(t, i) for i, t in enumerate(doc.tables, start=1)]

    full_text = "\n".join(p["text"] for p in paragraphs if p["text"])
    txt_path = output_dir / f"{doc_path.stem}_full_text.txt"
    txt_path.write_text(full_text, encoding="utf-8")

    return {
        "source": str(doc_path),
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        "engine": "python-docx",
        "properties": {"core": _core_properties(doc)},
        "paragraphs": paragraphs,
        "tables": tables,
        "statistics": {
            "paragraphs": len(paragraphs),
            "tables": len(tables),
        },
        "embedded_parts": _embedded_parts(doc, output_dir / "embedded_parts"),
        "full_text_export": str(txt_path),
    }
